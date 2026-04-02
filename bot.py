import os
import asyncio
import logging
import tempfile
from starlette.applications import Starlette
from starlette.routing import Route
from starlette.requests import Request
from starlette.responses import PlainTextResponse, Response
from telegram import Update, ReplyKeyboardMarkup, KeyboardButton
from telegram.ext import Application, CommandHandler, ConversationHandler, MessageHandler, filters, ContextTypes
from docx import Document
import uvicorn

logging.basicConfig(level=logging.INFO)
TOKEN = os.environ["BOT_TOKEN"]
URL = os.environ["RENDER_EXTERNAL_URL"]
PORT = int(os.getenv("PORT", 8000))

# ---------- ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ДЛЯ КЛАВИАТУР ----------
def make_keyboard(options, include_custom=True):
    """Создаёт клавиатуру из списка options, добавляет кнопку '✏️ Свой вариант' если include_custom=True"""
    if not options:
        return None
    buttons = [[KeyboardButton(opt)] for opt in options]
    if include_custom:
        buttons.append([KeyboardButton("✏️ Свой вариант")])
    return ReplyKeyboardMarkup(buttons, one_time_keyboard=True, resize_keyboard=True)

# ---------- КАЛЬКУЛЯТОР МНОАР ----------
# Состояния для калькулятора МНОАР
MNOAR_STATE_AGE, MNOAR_STATE_CARDIO, MNOAR_STATE_LUNG, MNOAR_STATE_OP = range(4)

async def mnoar_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Запускает калькулятор МНОАР"""
    await update.message.reply_text(
        "🧮 **Калькулятор МНОАР**\n\n"
        "Отвечайте на вопросы. Баллы суммируются автоматически.\n"
        "Для отмены введите /cancel_mnoar\n\n"
        "1️⃣ Возраст:\n"
        "- до 60 лет → 0 баллов\n"
        "- 60–70 лет → 1 балл\n"
        "- старше 70 лет → 2 балла\n\n"
        "Сколько лет пациенту?"
    )
    return MNOAR_STATE_AGE

async def mnoar_age(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        age = int(update.message.text.strip())
        if age < 60:
            points = 0
        elif age <= 70:
            points = 1
        else:
            points = 2
        context.user_data['mnoar_points'] = points
        await update.message.reply_text(
            f"✅ Возраст: {age} лет → {points} балл(ов).\n\n"
            "2️⃣ Сердечно-сосудистые заболевания:\n"
            "- нет → 0 баллов\n"
            "- компенсированные (гипертония I-II ст., стабильная стенокардия) → 1 балл\n"
            "- декомпенсированные (нестабильная стенокардия, ХСН, перенесённый инфаркт) → 3 балла\n\n"
            "Выберите вариант (можно написать свой):",
            reply_markup=make_keyboard(["нет", "компенсированные", "декомпенсированные"])
        )
        return MNOAR_STATE_CARDIO
    except ValueError:
        await update.message.reply_text("❌ Введите число (возраст целыми годами):")
        return MNOAR_STATE_AGE

async def mnoar_cardio(update: Update, context: ContextTypes.DEFAULT_TYPE):
    answer = update.message.text.strip().lower()
    if "компенсир" in answer:
        points = 1
    elif "декомпенсир" in answer:
        points = 3
    else:
        points = 0
    context.user_data['mnoar_points'] += points
    await update.message.reply_text(
        f"✅ Добавлено {points} балл(ов). Сумма: {context.user_data['mnoar_points']}\n\n"
        "3️⃣ Заболевания лёгких:\n"
        "- нет → 0 баллов\n"
        "- ХОБЛ, бронхиальная астма (компенсированные) → 1 балл\n"
        "- дыхательная недостаточность, зависимость от кислорода → 3 балла\n\n"
        "Выберите вариант:",
        reply_markup=make_keyboard(["нет", "ХОБЛ/астма", "дыхательная недостаточность"])
    )
    return MNOAR_STATE_LUNG

async def mnoar_lung(update: Update, context: ContextTypes.DEFAULT_TYPE):
    answer = update.message.text.strip().lower()
    if "хобл" in answer or "астма" in answer:
        points = 1
    elif "дыхательная недостаточность" in answer:
        points = 3
    else:
        points = 0
    context.user_data['mnoar_points'] += points
    await update.message.reply_text(
        f"✅ Добавлено {points} балл(ов). Сумма: {context.user_data['mnoar_points']}\n\n"
        "4️⃣ Характер операции:\n"
        "- малая (до 30 мин) → 0 баллов\n"
        "- средняя (30–120 мин) → 1 балл\n"
        "- большая (>120 мин) → 2 балла\n"
        "- экстренная (добавить 3 балла к основному риску) → 3 балла\n\n"
        "Выберите вариант:",
        reply_markup=make_keyboard(["малая", "средняя", "большая", "экстренная"])
    )
    return MNOAR_STATE_OP

async def mnoar_op(update: Update, context: ContextTypes.DEFAULT_TYPE):
    answer = update.message.text.strip().lower()
    if "малая" in answer:
        points = 0
    elif "средняя" in answer:
        points = 1
    elif "большая" in answer:
        points = 2
    elif "экстренная" in answer:
        points = 3
    else:
        points = 0
    total = context.user_data['mnoar_points'] + points
    # Оценка риска
    if total <= 2:
        risk = "низкий риск"
    elif total <= 5:
        risk = "средний риск"
    else:
        risk = "высокий риск"
    result = f"{total} баллов – {risk}"
    context.user_data['mnoar_result'] = result
    await update.message.reply_text(
        f"✅ Итоговый счёт МНОАР: {result}\n\n"
        "Это значение будет сохранено в поле {{mnoar}}.",
        reply_markup=ReplyKeyboardMarkup([], resize_keyboard=True)
    )
    # Завершаем калькулятор и возвращаем результат в основной диалог
    return ConversationHandler.END

async def mnoar_cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Калькулятор МНОАР отменён. Продолжаем основной опрос.", reply_markup=ReplyKeyboardMarkup([], resize_keyboard=True))
    return ConversationHandler.END


FIELDS = [
FIELDS = [
    # ----- Паспортная часть -----
    ("date", "📅 Дата", None, "text"),
    ("time", "⏰ Время", None, "text"),
    ("fio", "👤 ФИО пациента", None, "text"),
    ("age", "📆 Возраст (лет)", None, "number"),
    ("h", "📏 Рост (см)", None, "number"),
    ("w", "⚖️ Вес (кг)", None, "number"),
    # ИМТ рассчитается автоматически

    # ----- Порядок поступления и диагноз -----
    ("admission_order", "🏥 Порядок поступления", ["плановый", "срочный"], "choice_with_custom"),
    ("diagnosis", "📌 Диагноз (основной)", None, "text"),

    # ----- Жалобы -----
    ("complaints", "💬 Жалобы", ["активно нет", "болевой синдром в пояснице", "болевой синдром с иррадиацией в нижние конечности"], "choice_with_custom"),

    # ----- Анамнез заболевания -----
    ("history", "📋 Анамнез заболевания", ["дообследован амбулаторно, показания к операции", "госпитализирован срочно, обследован в соматически допустимом режиме"], "choice_with_custom"),

    # ----- Сопутствующие заболевания -----
    ("concomitant", "🩺 Сопутствующие заболевания", ["сахарный диабет 1 типа", "сахарный диабет 2 типа", "гипертоническая болезнь", "гипертиреоз", "гипотиреоз"], "choice_with_custom_multiple"),  # можно выбрать несколько

    # ----- Постоянный приём лекарств -----
    ("meds", "💊 Постоянный приём лекарств", None, "text"),

    # ----- Аллергологический анамнез -----
    ("allergy", "🌿 Аллергологический анамнез", ["нет", "есть (указать аллерген и проявление)"], "choice_with_custom"),

    # ----- Переливания крови -----
    ("blood", "🩸 Переливания крови", ["нет", "да, без осложнений", "да, с осложнениями (какими?)"], "choice_with_custom"),

    # ----- Нейроинфекции, ЧМТ -----
    ("neuro", "🧠 Нейроинфекции, ЧМТ", ["нет", "да"], "choice_with_custom"),

    # ----- Инфекционные заболевания -----
    ("inf", "🦠 Инфекционные заболевания (ВИЧ, гепатиты и др.)", None, "text"),

    # ----- Оперативные вмешательства -----
    ("ops", "🔪 Оперативные вмешательства в прошлом", ["нет", "да (какие?)"], "choice_with_custom"),

    # ----- Переносимость наркозов -----
    ("narc_tol", "💉 Переносимость наркозов", ["б/о", "тошнота", "головокружение", "долгое пробуждение"], "choice_with_custom_multiple"),

    # ----- Особенности -----
    ("specs", "📌 Особенности (дополнительно)", None, "text"),

    # ----- Объективный статус -----
    ("st_gen", "😐 Общее состояние", ["удовлетворительное", "средней тяжести", "тяжелое"], "choice_with_custom"),
    ("psycho", "🧠 Сознание", ["ясное", "спутанное", "оглушение"], "choice_with_custom"),
    ("body_type", "🏋️ Телосложение", ["астеник", "нормостеник", "гиперстеник"], "choice_with_custom"),
    ("obesity", "🍔 Ожирение (степень, если есть)", None, "text"),
    ("skin", "🩻 Кожные покровы и слизистые", ["обычной окраски", "бледные", "гиперемированы"], "choice_with_custom"),
    ("moist", "💧 Влажность кожи", ["нормальная", "снижена", "потливость"], "choice_with_custom"),
    ("thyroid", "🦋 Щитовидная железа", ["не увеличена", "увеличена"], "choice_with_custom"),
    ("turgor", "🤚 Тургор кожи", ["нормальный", "снижен"], "choice_with_custom"),
    ("temp", "🌡️ Температура тела", None, "number"),
    ("nodes", "🪢 Периферические лимфоузлы", ["не увеличены", "увеличены"], "choice_with_custom"),
    ("throat", "👄 Зев", ["не гиперемирован", "гиперемирован"], "choice_with_custom"),
    ("mallampaty", "😮 Mallampati", ["1", "2", "3", "4"], "choice_with_custom"),
    ("teeth", "🦷 Зубы и протезы", ["санированы", "не санированы, протезов нет", "есть съемные протезы"], "choice_with_custom"),
    ("edema", "💦 Периферические отёки", ["нет", "есть (где)"], "choice_with_custom"),
    ("breath", "🌬️ Дыхание", ["везикулярное", "бронхиальное"], "choice_with_custom"),
    ("rr", "📈 ЧДД (в мин)", None, "number"),
    ("wheeze", "🎵 Хрипы", ["нет", "сухие", "влажные", "сухие и влажные"], "choice_with_custom"),
    ("h_sounds", "❤️ Сердечные тоны", ["ясные, ритмичные", "приглушены, ритмичные", "глухие, аритмичные"], "choice_with_custom"),
    ("ps", "💓 Пульс (уд/мин)", None, "number"),
    ("p_def", "📉 Дефицит пульса", ["нет", "есть"], "choice_with_custom"),
    ("ad_s", "🩸 АД систолическое", None, "number"),
    ("ad_d", "🩸 АД диастолическое", None, "number"),
    ("spo2", "🫁 SpO2 (%)", None, "number"),
    ("tongue", "👅 Язык", ["влажный, не обложен", "сухой, обложен"], "choice_with_custom"),
    ("abd", "🤰 Живот", ["мягкий, безболезненный", "напряжённый, болезненный"], "choice_with_custom"),
    ("perist", "🔊 Перистальтика", ["выслушивается, нормальная", "ослаблена", "усилена"], "choice_with_custom"),
    ("liver", "🧬 Печень", ["не увеличена", "увеличена"], "choice_with_custom"),
    ("ur", "🚽 Мочеиспускание", ["свободное", "затруднённое"], "choice_with_custom"),
    ("stool", "💩 Стул", ["норма", "жидкий", "задержка (дней)"], "choice_with_custom"),
    ("hvn", "🦵 ХВН сосудов ног", ["нет", "есть (степень)"], "choice_with_custom"),
    ("lab", "🧪 Критические данные лаборатории", None, "text"),
    ("ecg", "📈 ЭКГ", None, "text"),

    # ----- ASA и МНОАР -----
    ("asa", "ASA", ["I", "II", "III", "IV", "V", "E"], "choice_with_custom"),
    ("mnoar", "📋 МНОАР (баллы)", None, "mnoar_calc"),  # специальный тип - вызов калькулятора

    # ----- Операция и анестезия -----
    ("op_vol", "📐 Объём операции", None, "text"),
    ("anes_type", "💉 Тип анестезии", ["тотальная", "сочетанная", "комбинированная", "в/венная", "ингаляционная", "эпидуральная", "субарахноидальная", "проводниковая"], "choice_with_custom"),
    ("add_test", "🔬 Дообследование", None, "text"),
    ("prep", "🛁 Предоперационная подготовка (кроме клизмы)", None, "text"),
    ("p_date", "📅 Дата премедикации", None, "text"),
    ("sib_dose", "💊 Сибазон (доза в мл)", None, "number"),
    ("time_before_op", "⏱️ За сколько минут до операции", None, "number"),
    ("prom_dose", "💉 Промедол 2% (доза в мл)", None, "number"),
    ("doc_name", "✍️ Фамилия врача", None, "text"),
]

]

STATE_LIST = list(range(len(FIELDS)))

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data.clear()
    context.user_data["step"] = 0
    field = FIELDS[0]
    question = field[1]
    options = field[2]
    input_type = field[3]
    include_custom = "choice" in input_type
    keyboard = make_keyboard(options, include_custom) if options else None
    await update.message.reply_text(
        "👨‍⚕️ **Заполнение осмотра анестезиолога**\n\nДля отмены /cancel\n\n" + question,
        parse_mode="Markdown",
        reply_markup=keyboard
    )
    return STATE_LIST[0]

async def handle_field(update: Update, context: ContextTypes.DEFAULT_TYPE):
    step = context.user_data.get("step", 0)
    if step >= len(FIELDS):
        return await generate_document(update, context)

    field_name, question, options, input_type = FIELDS[step]
    answer = update.message.text.strip()

    # Обработка кнопки "Свой вариант"
    if answer == "✏️ Свой вариант":
        await update.message.reply_text("Введите свой вариант текстом:")
        return STATE_LIST[step]  # остаёмся на том же шаге

    # Если поле с множественным выбором
    if input_type == "choice_with_custom_multiple":
        # Получаем уже выбранные ответы из user_data (если есть)
        multiple_key = f"multiple_{field_name}"
        if multiple_key not in context.user_data:
            context.user_data[multiple_key] = []
        if answer not in ["Готово", "закончить"]:
            context.user_data[multiple_key].append(answer)
            await update.message.reply_text(
                f"✅ Добавлено: {answer}\n\nМожно добавить ещё или напишите 'Готово' для продолжения.",
                reply_markup=make_keyboard(options + ["Готово"], include_custom=False)
            )
            return STATE_LIST[step]
        else:
            # Сохраняем собранный список в основное поле
            context.user_data[field_name] = ", ".join(context.user_data[multiple_key])
            del context.user_data[multiple_key]
            # Переходим к следующему шагу
            step += 1
            context.user_data["step"] = step
            if step >= len(FIELDS):
                await update.message.reply_text("✅ Все данные собраны. Формирую документ...")
                return await generate_document(update, context)
            else:
                next_field = FIELDS[step]
                next_question = next_field[1]
                next_options = next_field[2]
                next_input_type = next_field[3]
                include_custom = "choice" in next_input_type
                keyboard = make_keyboard(next_options, include_custom) if next_options else None
                await update.message.reply_text(next_question, reply_markup=keyboard)
                return STATE_LIST[step]

    # Если поле требует вызова калькулятора МНОАР
    if input_type == "mnoar_calc":
        context.user_data["pending_field"] = field_name
        await update.message.reply_text("Запускаю калькулятор МНОАР...")
        # Запускаем отдельный диалог калькулятора
        return await mnoar_start(update, context)

    # Обычная валидация
    if input_type == "number":
        try:
            float(answer)
        except ValueError:
            await update.message.reply_text("❌ Введите число (например, 70):")
            return STATE_LIST[step]
    elif "choice" in input_type and options and answer not in options and answer != "✏️ Свой вариант":
        await update.message.reply_text(f"❌ Пожалуйста, выберите вариант из кнопок или '✏️ Свой вариант': {', '.join(options)}")
        return STATE_LIST[step]

    # Сохраняем ответ
    context.user_data[field_name] = answer

    # Автоматический расчёт ИМТ после ввода веса
    if field_name == "w" and "h" in context.user_data and "bmi" not in context.user_data:
        try:
            h_cm = float(context.user_data["h"])
            w_kg = float(answer)
            h_m = h_cm / 100
            bmi = round(w_kg / (h_m ** 2), 1)
            context.user_data["bmi"] = str(bmi)
            logging.info(f"Рассчитан ИМТ: {bmi}")
        except:
            pass

    step += 1
    context.user_data["step"] = step

    if step >= len(FIELDS):
        await update.message.reply_text("✅ Все данные собраны. Формирую документ...")
        return await generate_document(update, context)
    else:
        next_field = FIELDS[step]
        next_question = next_field[1]
        next_options = next_field[2]
        next_input_type = next_field[3]
        include_custom = "choice" in next_input_type
        keyboard = make_keyboard(next_options, include_custom) if next_options else None
        await update.message.reply_text(next_question, reply_markup=keyboard)
        return STATE_LIST[step]

async def generate_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    data = context.user_data
    template_path = "template.docx"
    if not os.path.exists(template_path):
        await update.message.reply_text("❌ Ошибка: файл шаблона template.docx не найден.")
        return ConversationHandler.END

    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    with open(tmp_path, "rb") as f:
        await update.message.reply_document(
            document=f,
            filename=f"осмотр_{data.get('fio', 'patient')}.docx",
            caption="📄 Осмотр анестезиолога готов."
        )
    os.unlink(tmp_path)
    await update.message.reply_text("Для нового осмотра нажмите /start", reply_markup=ReplyKeyboardMarkup([], resize_keyboard=True))
    return ConversationHandler.END

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("❌ Заполнение отменено. /start для нового осмотра", reply_markup=ReplyKeyboardMarkup([], resize_keyboard=True))
    return ConversationHandler.END

# ---------- НАСТРОЙКА ВЕБХУКА И ЗАПУСК ----------
async def setup_webhook(application):
    webhook_url = f"{URL}/telegram"
    await application.bot.set_webhook(url=webhook_url, allowed_updates=Update.ALL_TYPES)
    logging.info(f"Webhook set to {webhook_url}")

async def main():
    application = Application.builder().token(TOKEN).updater(None).build()

    # Основной диалог
    conv_handler = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={state: [MessageHandler(filters.TEXT & ~filters.COMMAND, handle_field)] for state in STATE_LIST},
        fallbacks=[CommandHandler("cancel", cancel)],
    )
    application.add_handler(conv_handler)
    application.add_handler(CommandHandler("start", start))

    # Диалог калькулятора МНОАР (может вызываться из основного)
    mnoar_handler = ConversationHandler(
        entry_points=[CommandHandler("mnoar", mnoar_start)],
        states={
            MNOAR_STATE_AGE: [MessageHandler(filters.TEXT & ~filters.COMMAND, mnoar_age)],
            MNOAR_STATE_CARDIO: [MessageHandler(filters.TEXT & ~filters.COMMAND, mnoar_cardio)],
            MNOAR_STATE_LUNG: [MessageHandler(filters.TEXT & ~filters.COMMAND, mnoar_lung)],
            MNOAR_STATE_OP: [MessageHandler(filters.TEXT & ~filters.COMMAND, mnoar_op)],
        },
        fallbacks=[CommandHandler("cancel_mnoar", mnoar_cancel)],
    )
    application.add_handler(mnoar_handler)

    await setup_webhook(application)
    await application.initialize()
    await application.start()

    async def telegram_webhook(request: Request) -> Response:
        data = await request.json()
        update = Update.de_json(data, application.bot)
        await application.process_update(update)
        return Response()

    async def health_check(request: Request) -> PlainTextResponse:
        return PlainTextResponse("OK")

    starlette_app = Starlette(routes=[
        Route("/", health_check, methods=["GET"]),
        Route("/telegram", telegram_webhook, methods=["POST"]),
        Route("/healthcheck", health_check, methods=["GET"]),
    ])

    config = uvicorn.Config(starlette_app, host="0.0.0.0", port=PORT, log_level="info")
    server = uvicorn.Server(config)
    await server.serve()
    await application.stop()

if __name__ == "__main__":
    asyncio.run(main())